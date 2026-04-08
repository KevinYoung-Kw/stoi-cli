#!/usr/bin/env python3
"""
STOI Technical Design Document Generator
使用 python-docx 生成专业的Word文档
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_TABLE_ALIGNMENT

def set_cell_shading(cell, color):
    """设置单元格背景色"""
    from docx.oxml.ns import qn
    from docx.oxml import parse_xml
    from docx.oxml import OxmlElement

    tcPr = cell._tc.get_or_add_tcPr()
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color)
    tcPr.append(shading)

def main():
    doc = Document()

    # 设置默认字体
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(11)

    # 设置标题样式
    for i, (name, size, color) in enumerate([
        ('Heading 1', 18, '2E74B5'),
        ('Heading 2', 14, '2E74B5'),
        ('Heading 3', 12, '5B9BD5'),
    ], 1):
        style = doc.styles[name]
        font = style.font
        font.name = 'Arial'
        font.size = Pt(size)
        font.bold = True
        font.color.rgb = RGBColor.from_string(color)

    # ===== 封面 =====
    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph()

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run('STOI')
    run.bold = True
    run.font.size = Pt(48)
    run.font.color.rgb = RGBColor.from_string('2E74B5')

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run('Shit Token On Investment')
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor.from_string('666666')

    doc.add_paragraph()

    title2 = doc.add_paragraph()
    title2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title2.add_run('技术设计文档')
    run.bold = True
    run.font.size = Pt(28)
    run.font.color.rgb = RGBColor.from_string('2E74B5')

    doc.add_paragraph()
    doc.add_paragraph()

    info = doc.add_paragraph()
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    info.add_run('Technical Design Document\n\n').font.size = Pt(14)
    info.add_run('版本: v1.0\n').font.size = Pt(12)
    info.add_run('日期: 2026-04-08\n').font.size = Pt(12)
    info.add_run('文档状态: 技术评审').font.size = Pt(12)

    doc.add_page_break()

    # ===== 目录 =====
    doc.add_heading('目录', level=1)
    toc_items = [
        '1. 执行摘要 ............................................................. 3',
        '2. 产品概述 ............................................................. 4',
        '3. 技术架构 ............................................................. 5',
        '4. 核心模块设计 ......................................................... 7',
        '5. 数据模型 ............................................................. 12',
        '6. CLI 命令设计 ......................................................... 15',
        '7. 开发路线图 ........................................................... 18',
        '8. 技术风险与缓解 ....................................................... 19',
        '9. 参考来源 ............................................................. 20',
    ]
    for item in toc_items:
        doc.add_paragraph(item, style='Normal')

    doc.add_page_break()

    # ===== 1. 执行摘要 =====
    doc.add_heading('1. 执行摘要', level=1)

    p = doc.add_paragraph()
    p.add_run('STOI (Shit Token On Investment) ').bold = True
    p.add_run('是一款专为 AI 编程工具设计的 CLI 效率监控工具。它通过量化 Token 消耗的转化率，帮助开发者和企业识别"有效代码输出 (Value)"与"无效/冗余消耗 (Shit)"的比例，最终给出直观的「含屎量」评级。')

    doc.add_heading('1.1 核心价值主张', level=2)
    bullets = [
        ('成本控制: ', '识别并减少因 Prompt 设计缺陷导致的 Token 浪费'),
        ('效率提升: ', '量化 AI 辅助编程的真实 ROI'),
        ('生态优化: ', '通过生成 Issue 报告倒逼框架改进'),
    ]
    for title, content in bullets:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(title).bold = True
        p.add_run(content)

    doc.add_heading('1.2 技术选型总览', level=2)

    # 创建表格
    table = doc.add_table(rows=6, cols=2)
    table.style = 'Table Grid'

    # 表头
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = '层级'
    hdr_cells[1].text = '技术选型'
    set_cell_shading(hdr_cells[0], '2E74B5')
    set_cell_shading(hdr_cells[1], '2E74B5')
    for cell in hdr_cells:
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
        cell.paragraphs[0].runs[0].bold = True

    # 数据行
    data = [
        ('CLI 框架', 'Typer + Rich (Python)'),
        ('数据存储', 'SQLite + Pydantic'),
        ('LLM 评估', 'GPT-4.1 + Claude-4 + Gemini-2'),
        ('拦截机制', 'Claude Code Hooks + MCP 中间件'),
        ('Diff 算法', '语义 + 行级 + Token 级多层 Diff'),
    ]
    for i, (col1, col2) in enumerate(data, 1):
        row = table.rows[i].cells
        row[0].text = col1
        row[1].text = col2

    doc.add_page_break()

    # ===== 2. 产品概述 =====
    doc.add_heading('2. 产品概述', level=1)

    doc.add_heading('2.1 背景与痛点', level=2)
    doc.add_paragraph(
        '随着 Claude Code 等 AI 编程工具的普及，开发者发现系统提示词中的硬编码时间戳导致 KV-Cache 大面积不命中，'
        '引发 Token 的无谓暴涨（Token 刺客现象）。然而，业界目前只有 Token 消耗计费，却缺乏衡量 Token 产出比（TOI）的指标。'
    )

    doc.add_heading('2.2 核心痛点', level=2)
    pain_points = [
        ('缺乏衡量标准: ', '开发者不知道消耗的 Token 是转化成了高价值代码，还是变成了毫无意义的废话'),
        ('企业效能黑盒: ', '难以界定员工 AI 辅助编程的真实效率'),
        ('成本失控: ', '每次无效的上下文对话都在"烧钱"'),
    ]
    for i, (title, content) in enumerate(pain_points, 1):
        p = doc.add_paragraph(style='List Number')
        p.add_run(title).bold = True
        p.add_run(content)

    doc.add_heading('2.3 产品定位', level=2)
    doc.add_paragraph(
        'STOI 是开发者侧首个基于 CLI 形态的「词元投资回报率」分析与监控工具。它通过监控底层 Token 消耗，'
        '结合 LLM-as-a-judge 机制，为开发者和企业量化 Token 效率，最终给出直观的「含屎量」评级和改进建议。'
    )

    doc.add_page_break()

    # ===== 3. 技术架构 =====
    doc.add_heading('3. 技术架构', level=1)

    doc.add_heading('3.1 架构全景', level=2)
    doc.add_paragraph(
        'STOI 采用分层架构设计，包含 CLI 层、核心服务层、数据采集层、评估层和存储层。'
        '各层之间通过明确定义的接口进行交互，确保系统的可扩展性和可维护性。'
    )

    doc.add_heading('3.2 核心组件', level=2)

    table = doc.add_table(rows=5, cols=3)
    table.style = 'Table Grid'

    hdr = table.rows[0].cells
    hdr[0].text, hdr[1].text, hdr[2].text = '组件', '职责', '技术选型'
    for cell in hdr:
        set_cell_shading(cell, 'D9E2F3')
        cell.paragraphs[0].runs[0].bold = True

    components = [
        ('Session Manager', '会话生命周期管理', 'Python asyncio'),
        ('Token Analyzer', 'Token 效率分析', '多模型集成评估'),
        ('Cache Inspector', 'KV-Cache 分析', 'Diff 算法 + 模式匹配'),
        ('Report Generator', '报告生成与导出', 'Jinja2 + Markdown'),
    ]
    for i, (name, duty, tech) in enumerate(components, 1):
        row = table.rows[i].cells
        row[0].text = name
        row[1].text = duty
        row[2].text = tech

    doc.add_heading('3.3 数据流', level=2)
    data_flow = [
        '数据采集: Hook 拦截 → MCP 中间件 → JSONL 解析',
        '实时处理: 数据清洗 → 模式检测 → Cache 事件标记',
        '深度分析: LLM 评估 → 冗余检测 → 价值提取',
        '存储索引: SQLite 写入 → 聚合统计 → 报表生成',
    ]
    for item in data_flow:
        doc.add_paragraph(item, style='List Number')

    doc.add_page_break()

    # ===== 4. 核心模块设计 =====
    doc.add_heading('4. 核心模块设计', level=1)

    doc.add_heading('4.1 LLM-as-Judge 评估体系', level=2)
    doc.add_paragraph('评估维度与权重').runs[0].bold = True

    table = doc.add_table(rows=5, cols=4)
    table.style = 'Table Grid'

    hdr = table.rows[0].cells
    hdr[0].text, hdr[1].text, hdr[2].text, hdr[3].text = '维度', '权重', '评估内容', '评分方法'
    for cell in hdr:
        set_cell_shading(cell, 'D9E2F3')
        cell.paragraphs[0].runs[0].bold = True

    dimensions = [
        ('问题解决度', '35%', '功能完整性、边界处理、实际可用性', '1-5 分 Likert 量表'),
        ('代码质量', '25%', '正确性、可读性、最佳实践、健壮性', '多因子加权'),
        ('信息密度', '20%', '信噪比、重复率、相关度', 'Token 级分析'),
        ('上下文效率', '20%', '利用率、精准度、选择性', '上下文追踪'),
    ]
    for i, (dim, weight, content, method) in enumerate(dimensions, 1):
        row = table.rows[i].cells
        row[0].text = dim
        row[1].text = weight
        row[2].text = content
        row[3].text = method

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run('STOI 指数计算公式').bold = True
    doc.add_paragraph('STOI_Index = Σ(Dimension_Score_i × Weight_i) × Efficiency_Factor × Reliability_Factor')
    doc.add_paragraph('含屎量 (Waste Ratio) = 1 - (Effective_Tokens / Total_Tokens)')

    doc.add_heading('4.2 Token 拦截与 Cache 分析', level=2)
    doc.add_paragraph('多层拦截架构').runs[0].bold = True

    layers = [
        ('Layer 1: ', 'Claude Code Hooks (PreToolUse/PostToolUse) - 事件捕获'),
        ('Layer 2: ', 'MCP Middleware - 请求/响应拦截'),
        ('Layer 3: ', 'Statusbar Monitor - Token 计数轮询'),
        ('Layer 4: ', 'JSONL Log Parser - 异步分析'),
    ]
    for title, content in layers:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(title).bold = True
        p.add_run(content)

    doc.add_paragraph()
    doc.add_paragraph('Cache Miss 检测模式').runs[0].bold = True

    table = doc.add_table(rows=4, cols=3)
    table.style = 'Table Grid'

    hdr = table.rows[0].cells
    hdr[0].text, hdr[1].text, hdr[2].text = '模式类型', '描述', '影响级别'
    for cell in hdr:
        set_cell_shading(cell, 'D9E2F3')
        cell.paragraphs[0].runs[0].bold = True

    patterns = [
        ('ISO_TIMESTAMP', 'ISO 8601 时间戳格式', 'CRITICAL'),
        ('UUID_V4', 'UUID 版本 4 格式', 'CRITICAL'),
        ('ABS_PATH', 'Unix/Windows 绝对路径', 'HIGH'),
    ]
    for i, (pattern, desc, impact) in enumerate(patterns, 1):
        row = table.rows[i].cells
        row[0].text = pattern
        row[1].text = desc
        row[2].text = impact

    doc.add_page_break()

    # ===== 5. 数据模型 =====
    doc.add_heading('5. 数据模型', level=1)
    doc.add_paragraph(
        'STOI 采用关系型数据模型，核心实体包括 Session、Message、TokenUsage、Evaluation 和 CacheEvent。'
    )

    doc.add_heading('5.1 Session（会话）', level=2)
    table = doc.add_table(rows=5, cols=3)
    table.style = 'Table Grid'

    hdr = table.rows[0].cells
    hdr[0].text, hdr[1].text, hdr[2].text = '字段', '类型', '说明'
    for cell in hdr:
        set_cell_shading(cell, 'D9E2F3')
        cell.paragraphs[0].runs[0].bold = True

    session_fields = [
        ('id', 'UUID', '主键'),
        ('name', 'String', '会话名称'),
        ('project', 'String', '项目标识'),
        ('start_time', 'DateTime', '开始时间'),
    ]
    for i, (field, type_, desc) in enumerate(session_fields, 1):
        row = table.rows[i].cells
        row[0].text = field
        row[1].text = type_
        row[2].text = desc

    doc.add_heading('5.2 TokenUsage（Token 使用）', level=2)
    table = doc.add_table(rows=5, cols=3)
    table.style = 'Table Grid'

    hdr = table.rows[0].cells
    hdr[0].text, hdr[1].text, hdr[2].text = '字段', '类型', '说明'
    for cell in hdr:
        set_cell_shading(cell, 'D9E2F3')
        cell.paragraphs[0].runs[0].bold = True

    token_fields = [
        ('message_id', 'UUID', '关联消息'),
        ('input_tokens', 'Integer', '输入 Token 数'),
        ('cache_read_tokens', 'Integer', 'Cache 读取 Token'),
        ('cache_creation_tokens', 'Integer', 'Cache 创建 Token'),
    ]
    for i, (field, type_, desc) in enumerate(token_fields, 1):
        row = table.rows[i].cells
        row[0].text = field
        row[1].text = type_
        row[2].text = desc

    doc.add_page_break()

    # ===== 6. CLI 命令设计 =====
    doc.add_heading('6. CLI 命令设计', level=1)

    doc.add_heading('6.1 核心命令清单', level=2)
    table = doc.add_table(rows=6, cols=3)
    table.style = 'Table Grid'

    hdr = table.rows[0].cells
    hdr[0].text, hdr[1].text, hdr[2].text = '命令', '功能', '关键参数'
    for cell in hdr:
        set_cell_shading(cell, '2E74B5')
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
        cell.paragraphs[0].runs[0].bold = True

    commands = [
        ('stoi track', '开始/停止追踪', 'start/stop, --name'),
        ('stoi analyze', '效率分析', '--session, --format'),
        ('stoi blame', 'Cache 归因', '--pattern, --suggest'),
        ('stoi report', '生成报告', '--period, --output'),
        ('stoi config', '配置管理', 'show/edit/set'),
    ]
    for i, (cmd, func, params) in enumerate(commands, 1):
        row = table.rows[i].cells
        row[0].text = cmd
        row[0].paragraphs[0].runs[0].bold = True
        row[1].text = func
        row[2].text = params

    doc.add_heading('6.2 使用示例', level=2)

    examples = [
        ('开始追踪会话', 'stoi track start --name "API Refactor" --tags backend,python'),
        ('分析并导出报告', 'stoi analyze --format json --output analysis.json'),
        ('生成 Git Blame 式报告', 'stoi blame --suggest --threshold 60'),
    ]
    for title, cmd in examples:
        p = doc.add_paragraph()
        p.add_run(title).bold = True
        p = doc.add_paragraph()
        p.add_run(f'$ {cmd}').font.name = 'Courier New'
        p.paragraph_format.left_indent = Inches(0.25)

    doc.add_page_break()

    # ===== 7. 开发路线图 =====
    doc.add_heading('7. 开发路线图', level=1)

    phases = [
        ('Phase 1: MVP (Week 1-4)', [
            'CLI 框架搭建（Typer + Rich）',
            '数据模型实现（Pydantic + SQLite）',
            'stoi track / analyze 基础命令',
            'Hook 拦截器实现',
        ]),
        ('Phase 2: 分析引擎 (Week 5-8)', [
            'LLM-as-judge 评估模块',
            'Cache Miss 检测算法',
            'stoi blame 命令',
            '报告生成器',
        ]),
        ('Phase 3: 集成与优化 (Week 9-12)', [
            'MCP 服务器暴露',
            'Claude Code 深度集成',
            '性能优化与缓存',
        ]),
    ]

    for phase_title, tasks in phases:
        doc.add_heading(phase_title, level=2)
        for task in tasks:
            doc.add_paragraph(task, style='List Bullet')

    doc.add_page_break()

    # ===== 8. 技术风险与缓解 =====
    doc.add_heading('8. 技术风险与缓解', level=1)

    table = doc.add_table(rows=4, cols=3)
    table.style = 'Table Grid'

    hdr = table.rows[0].cells
    hdr[0].text, hdr[1].text, hdr[2].text = '风险', '影响', '缓解策略'
    for cell in hdr:
        set_cell_shading(cell, 'D9E2F3')
        cell.paragraphs[0].runs[0].bold = True

    risks = [
        ('Claude Code API 限制', '高', '多层拦截架构、Hook 系统、Statusbar 轮询'),
        ('LLM 评判偏见', '中', '多模型集成、偏见检测、自适应评判'),
        ('Token 计数不准确', '中', '使用 Statusbar API、JSONL 校验'),
    ]
    for i, (risk, impact, mitigation) in enumerate(risks, 1):
        row = table.rows[i].cells
        row[0].text = risk
        row[1].text = impact
        row[2].text = mitigation

    doc.add_page_break()

    # ===== 9. 参考来源 =====
    doc.add_heading('9. 参考来源', level=1)

    doc.add_heading('9.1 学术论文', level=2)
    papers = [
        'Bai et al. (2022). Constitutional AI: Harmlessness from AI Feedback. arXiv:2212.08073',
        'Zheng et al. (2023). Judging LLM-as-a-Judge with MT-Bench and Chatbot Arena. NeurIPS 2023',
        'Chen et al. (2021). Evaluating Large Language Models Trained on Code. arXiv:2107.03374',
        'MultiChallenge (2025). A Realistic Multi-Turn Conversation Evaluation Benchmark. arXiv:2501.17399',
    ]
    for i, paper in enumerate(papers, 1):
        doc.add_paragraph(f'{i}. {paper}')

    doc.add_heading('9.2 技术文档', level=2)
    docs = [
        'Claude Code Hooks Documentation',
        'MCP Protocol Specification',
        'Typer & Rich Official Documentation',
    ]
    for i, doc_item in enumerate(docs, 1):
        doc.add_paragraph(f'{i}. {doc_item}')

    doc.add_paragraph()
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run('--- 文档结束 ---').font.color.rgb = RGBColor.from_string('666666')

    # 保存文档
    output_path = '/Users/kevinyoung/Desktop/red-hackathon/STOI-Technical-Design-v1.0.docx'
    doc.save(output_path)
    print(f'✅ Document generated: {output_path}')

if __name__ == '__main__':
    main()
